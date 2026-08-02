from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from markdown_to_docx import (
    Block,
    HEADING_RE,
    add_inline_runs,
    parse_blocks,
)


BLUE = "2E74B5"
DARK_BLUE = "17365D"
MID_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F3F7FB"
GRAY = "666666"
LIGHT_GRAY = "F2F2F2"
USABLE_DXA = 9360


def set_font(run, name: str, size: float, color: str | None = None, bold: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:shd"))
    if old is not None:
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("Trang ")
    set_font(run, "Calibri", 9, GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])


def add_toc(document: Document, blocks: list[Block]) -> None:
    p = document.add_paragraph()
    p.style = document.styles["Heading 1"]
    p.add_run("Mục lục")
    for block in blocks:
        if block.kind != "heading":
            continue
        match = HEADING_RE.match(block.lines[0])
        if not match or len(match.group(1)) != 2:
            continue
        entry = document.add_paragraph(style="List Number")
        entry.paragraph_format.space_after = Pt(2)
        add_inline_runs(entry, match.group(2).strip(), font_size=10.5)
    document.add_page_break()


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, MID_BLUE, 10, 5),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.page_break_before = False

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

def add_cover(document: Document, title: str, subtitle: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(88)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("BALII SLEEPWEAR")
    set_font(r, "Calibri", 12, BLUE, True)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(title)
    set_font(r, "Calibri", 28, DARK_BLUE, True)

    line = document.add_paragraph()
    line.paragraph_format.space_after = Pt(20)
    line_pr = line._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    line_pr.append(borders)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run(subtitle)
    set_font(r, "Calibri", 14, MID_BLUE)

    box = document.add_paragraph()
    box.paragraph_format.left_indent = Inches(0.18)
    box.paragraph_format.right_indent = Inches(0.18)
    box.paragraph_format.space_before = Pt(12)
    box.paragraph_format.space_after = Pt(18)
    p_pr = box._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHTER_BLUE)
    p_pr.append(shd)
    r = box.add_run(
        "Bản viết lại dựa trên mã nguồn, migration, BPMN và cấu hình triển khai hiện tại. "
        "Tài liệu phân biệt rõ phần đã triển khai, phần có mô phỏng/fallback và phần còn thiếu."
    )
    set_font(r, "Calibri", 11, DARK_BLUE)

    meta = document.add_paragraph()
    meta.paragraph_format.space_before = Pt(72)
    meta.add_run("Phiên bản khảo sát: 20/07/2026\n")
    meta.add_run("Phạm vi: frontend, backend services, dữ liệu, workflow, event, AI và vận hành\n")
    meta.add_run("Ngôn ngữ: Tiếng Việt")
    for run in meta.runs:
        set_font(run, "Calibri", 10.5, GRAY)

    document.add_page_break()


def add_running_header_footer(document: Document) -> None:
    section = document.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Balii SleepWear  |  Tài liệu công nghệ hiện trạng")
    set_font(r, "Calibri", 9, GRAY)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(p)


def normalize_table_line(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def width_plan(rows: list[list[str]], total: int) -> list[int]:
    col_count = max(len(row) for row in rows)
    weights = [8] * col_count
    for row in rows:
        for i, value in enumerate(row):
            weights[i] = max(weights[i], min(42, len(value)))
    weight_sum = sum(weights)
    result = [round(total * weight / weight_sum) for weight in weights]
    result[-1] += total - sum(result)
    return result


def render_table(document: Document, lines: list[str]) -> None:
    rows = [normalize_table_line(line) for line in lines]
    rows = [rows[0], *rows[2:]]
    column_count = len(rows[0])
    widths = width_plan(rows, USABLE_DXA)
    table = document.add_table(rows=1, cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    set_cell_margins(table)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, values in enumerate(rows):
        cells = table.rows[0].cells if row_index == 0 else table.add_row().cells
        if row_index == 0:
            set_repeat_table_header(table.rows[0])
        for col_index, cell in enumerate(cells):
            set_cell_width(cell, widths[col_index])
            cell.width = Inches(widths[col_index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            text = values[col_index] if col_index < len(values) else ""
            add_inline_runs(p, text, font_name="Calibri", font_size=9.5)
            if row_index == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, LIGHT_BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def render_code(document: Document, lines: list[str]) -> None:
    for index, line in enumerate(lines):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.right_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(4 if index == 0 else 0)
        p.paragraph_format.space_after = Pt(4 if index == len(lines) - 1 else 0)
        p.paragraph_format.keep_together = True
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), LIGHT_GRAY)
        p_pr.append(shd)
        r = p.add_run(line or " ")
        set_font(r, "Consolas", 9, DARK_BLUE)


def render_blocks(document: Document, blocks: list[Block]) -> None:
    for block in blocks:
        if block.kind == "heading":
            match = HEADING_RE.match(block.lines[0])
            if match:
                level = min(3, len(match.group(1)))
                heading_text = match.group(2).strip()
                heading = document.add_paragraph(heading_text, style=f"Heading {level}")
                heading.paragraph_format.keep_with_next = True
                heading.paragraph_format.keep_together = True
                if heading_text in {
                    "3. API Gateway, xác thực và mô hình tin cậy",
                    "6.2 Transactional outbox",
                }:
                    heading.paragraph_format.page_break_before = True
        elif block.kind == "paragraph":
            p = document.add_paragraph()
            add_inline_runs(p, " ".join(block.lines))
        elif block.kind in ("bullet_list", "number_list"):
            numbered = block.kind == "number_list"
            style = "List Number" if numbered else "List Bullet"
            import re
            for line in block.lines:
                text = re.sub(r"^\d+\.\s+", "", line) if numbered else re.sub(r"^[-*+]\s+", "", line)
                p = document.add_paragraph(style=style)
                add_inline_runs(p, text)
        elif block.kind == "code":
            render_code(document, block.lines)
        elif block.kind == "table":
            render_table(document, block.lines)


def build(markdown_path: Path, output_path: Path) -> None:
    text = markdown_path.read_text(encoding="utf-8")
    blocks = parse_blocks(text)
    if len(blocks) < 2 or blocks[0].kind != "heading" or blocks[1].kind != "paragraph":
        raise ValueError("Markdown must start with a title and subtitle paragraph")

    title_match = HEADING_RE.match(blocks[0].lines[0])
    assert title_match is not None
    title = title_match.group(2).strip()
    subtitle = " ".join(blocks[1].lines)

    document = Document()
    configure(document)
    add_running_header_footer(document)
    add_cover(document, title, subtitle)
    add_toc(document, blocks[2:])
    render_blocks(document, blocks[2:])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.core_properties.title = title
    document.core_properties.subject = "Kiến trúc và công nghệ dự án Balii SleepWear"
    document.core_properties.author = "Balii SleepWear"
    document.core_properties.keywords = "Balii, NestJS, Next.js, PostgreSQL, Camunda, Kafka, RAG, Virtual Try-on"
    document.save(output_path)


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    source = root / "docs" / "deep-tech-explainer.vi.updated.md"
    target = root / "docs" / "deep-tech-explainer.vi.updated.docx"
    if len(sys.argv) > 1:
        source = Path(sys.argv[1]).resolve()
    if len(sys.argv) > 2:
        target = Path(sys.argv[2]).resolve()
    build(source, target)
    print(target)
