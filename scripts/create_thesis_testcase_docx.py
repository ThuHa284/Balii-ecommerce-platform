from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path("docs/Bo_test_case_luan_van_rut_gon.docx")
FONT = "Times New Roman"
NAVY = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"


FUNCTIONAL = [
    ("TC01", "Đăng ký", "Người dùng nhập email chưa tồn tại và mật khẩu hợp lệ.",
     "email: an.nguyen@test.com; password: Abc@123; fullName: Nguyễn Văn An",
     "Tài khoản được tạo; mật khẩu được mã hóa; hệ thống thông báo thành công."),
    ("TC02", "Đăng ký", "Người dùng nhập email đã tồn tại.",
     "email: user@test.com (đã có trong hệ thống); password: Abc@123",
     "Hệ thống từ chối, báo email đã được sử dụng và không tạo dữ liệu trùng."),
    ("TC03", "Đăng nhập", "Người dùng nhập đúng email và mật khẩu.",
     "email: user@test.com; password: Abc@123",
     "Hệ thống xác thực thành công, trả token và thông tin người dùng."),
    ("TC04", "Đăng nhập", "Người dùng nhập đúng email nhưng sai mật khẩu.",
     "email: user@test.com; password: Wrong123",
     "Hệ thống báo lỗi xác thực và không cấp token."),
    ("TC05", "Sản phẩm", "Người dùng mở danh sách sản phẩm.",
     "Không có tham số hoặc tham số phân trang hợp lệ",
     "Danh sách sản phẩm đang hoạt động được hiển thị đúng và có phân trang."),
    ("TC06", "Sản phẩm", "Admin thêm sản phẩm và biến thể có SKU mới.",
     "Tên, slug, SKU mới, giá > 0, tồn kho ≥ 0 và ảnh hợp lệ",
     "Sản phẩm, biến thể và ảnh được lưu thành công; dữ liệu hiển thị đúng."),
    ("TC07", "Giỏ hàng", "Khách vãng lai thêm một biến thể còn hàng vào giỏ.",
     "sessionId: guest-001; variantId hợp lệ; quantity: 1",
     "Mục hàng được thêm; giỏ được lưu theo cart:guest:{sessionId}; tổng tiền được tính lại."),
    ("TC08", "Giỏ hàng", "Người dùng đăng nhập sau khi đã có giỏ khách.",
     "Giỏ khách và giỏ người dùng đều có sản phẩm",
     "Giỏ khách được hợp nhất vào giỏ người dùng; không mất hoặc tạo dòng trùng sai."),
    ("TC09", "Checkout", "Người dùng đặt hàng với giỏ và địa chỉ hợp lệ.",
     "Giỏ còn hàng; địa chỉ đầy đủ; paymentMethod: cod",
     "Đơn hàng được tạo ở trạng thái PENDING; chi tiết và tổng tiền chính xác."),
    ("TC10", "Thanh toán", "Cổng thanh toán trả kết quả thành công với chữ ký hợp lệ.",
     "Mã phản hồi thành công; checksum hợp lệ; số tiền và mã đơn khớp",
     "Payment chuyển SUCCESS; Order được cập nhật trạng thái phù hợp; chỉ xử lý một lần."),
    ("TC11", "Thanh toán", "Cổng thanh toán trả kết quả thất bại hoặc người dùng hủy.",
     "Mã phản hồi thất bại/hủy; checksum hợp lệ",
     "Payment chuyển FAILED/CANCELLED; Order không được xác nhận đã thanh toán."),
    ("TC12", "Voucher", "Người dùng nhập voucher còn hạn và đủ điều kiện.",
     "code: SALE10; orderAmount: 500.000 đồng",
     "Voucher được chấp nhận; tổng tiền giảm đúng cấu hình và không vượt mức giảm tối đa."),
    ("TC13", "Try-On", "Người dùng đăng nhập, tải ảnh hợp lệ và chọn sản phẩm.",
     "Ảnh JPG/PNG hợp lệ; sản phẩm hợp lệ; mode: balanced",
     "Job thử đồ được tạo; trạng thái được theo dõi; trả resultUrl khi hoàn tất và lưu lịch sử."),
    ("TC14", "Market Analysis", "Admin yêu cầu phân tích giá theo từ khóa hợp lệ.",
     "keyword: váy công sở; limit: 10; saveResults: true",
     "Hệ thống thu thập, chuẩn hóa, loại trùng, lưu dữ liệu và sinh thông tin tổng hợp/insight."),
]


NON_FUNCTIONAL = [
    ("NFTC01", "Hiệu năng", "Đo thời gian phản hồi API danh sách sản phẩm.",
     "100 người dùng đồng thời trong 5 phút.",
     "P95 ≤ 2 giây; tỷ lệ lỗi < 1%."),
    ("NFTC02", "Hiệu năng", "Đo thời gian thêm giỏ và tạo đơn.",
     "50 người dùng đồng thời trong 5 phút.",
     "P95 thêm giỏ ≤ 1,5 giây; P95 tạo đơn ≤ 3 giây; lỗi < 1%."),
    ("NFTC03", "Bảo mật", "Truy cập API bảo vệ khi thiếu, sai hoặc hết hạn token.",
     "Không token; token bị sửa; token hết hạn.",
     "Trả 401; không trả dữ liệu nhạy cảm."),
    ("NFTC04", "Phân quyền", "Customer truy cập API quản trị hoặc dữ liệu của người khác.",
     "Token CUSTOMER; endpoint admin/mã đơn của tài khoản khác.",
     "Trả 403/404; dữ liệu không bị đọc hoặc sửa trái phép."),
    ("NFTC05", "Bảo mật đầu vào", "Gửi chuỗi SQL injection và XSS vào các trường nhập.",
     "' OR 1=1--; <script>alert(1)</script>.",
     "Không thực thi mã/lệnh; không lộ lỗi cơ sở dữ liệu."),
    ("NFTC06", "Tương thích", "Kiểm tra giao diện trên trình duyệt và màn hình phổ biến.",
     "Chrome, Edge, Firefox; 375×667, 768×1024, 1920×1080.",
     "Chức năng chính hoạt động; giao diện không vỡ hoặc che nút thao tác."),
    ("NFTC07", "Ổn định", "Chạy tải trung bình liên tục để phát hiện rò rỉ tài nguyên.",
     "50 người dùng đồng thời trong 2 giờ.",
     "CPU/RAM không tăng bất thường; tỷ lệ lỗi < 1%."),
    ("NFTC08", "Toàn vẹn dữ liệu", "Gửi đồng thời nhiều yêu cầu mua sản phẩm chỉ còn một đơn vị.",
     "Tồn kho: 1; 10 yêu cầu checkout đồng thời.",
     "Không bán vượt tồn kho; tối đa một yêu cầu thành công."),
]


EXCEPTIONS = [
    ("Xác thực", "Token thiếu, sai hoặc hết hạn", "Guard chặn request trước khi vào nghiệp vụ.", "Trả 401; không lộ dữ liệu."),
    ("Dữ liệu đầu vào", "Thiếu trường, sai kiểu hoặc chuỗi quá dài", "Validation từ chối và trả lỗi có kiểm soát.", "Trả 400; không tạo dữ liệu sai."),
    ("Sản phẩm", "Sản phẩm/biến thể bị vô hiệu hóa", "Kiểm tra trạng thái trước khi hiển thị và thêm giỏ.", "Không cho mua; thông báo sản phẩm không khả dụng."),
    ("Giỏ hàng", "Số lượng yêu cầu vượt tồn kho", "Đối chiếu tồn kho hiện tại.", "Yêu cầu điều chỉnh; không tạo đơn vượt kho."),
    ("Giỏ hàng", "Redis tạm thời không khả dụng", "Timeout sớm, ghi log và trả thông báo thử lại.", "API không treo và không báo lưu thành công sai."),
    ("Voucher", "Voucher hết hạn khi đang checkout", "Kiểm tra lại hiệu lực trong lúc tạo đơn.", "Không áp dụng giảm giá cũ; tính lại tổng tiền."),
    ("Checkout", "Người dùng nhấn đặt hàng nhiều lần", "Dùng khóa nghiệp vụ/idempotency và khóa nút khi đang gửi.", "Chỉ một đơn được tạo."),
    ("Thanh toán", "Callback sai checksum, sai số tiền hoặc bị gửi lặp", "Xác minh chữ ký và xử lý idempotent.", "Không cập nhật sai hoặc ghi nhận giao dịch hai lần."),
    ("Try-On", "Dịch vụ AI timeout hoặc trả lỗi", "Job chuyển trạng thái lỗi; retry giới hạn; ghi log.", "Giao diện không treo; người dùng có thể thử lại."),
    ("Market Analysis", "API crawl hết quota hoặc không phản hồi", "Timeout, retry có backoff và không đánh dấu thành công.", "Admin nhận thông báo phù hợp và có thể thử lại."),
    ("Cơ sở dữ liệu", "Mất kết nối hoặc deadlock", "Rollback transaction; retry có giới hạn.", "Không ghi dữ liệu một phần; có log truy vết."),
    ("Toàn hệ thống", "Lỗi không dự kiến", "Global exception filter/error boundary xử lý và ghi log.", "Không lộ stack trace hoặc secret; có thông báo thân thiện."),
]


def set_font(run, size=9, bold=False, color="000000"):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_twips = round(sum(widths_cm) / 2.54 * 1440)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_cm in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(round(width_cm / 2.54 * 1440)))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            twips = round(widths_cm[index] / 2.54 * 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(twips))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Cm(widths_cm[index])


def fill_cell(cell, value, header=False, center=False, font_size=8.7):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(value)
    set_font(run, size=font_size, bold=header, color=WHITE if header else "000000")
    if header:
        shade(cell, NAVY)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_data_table(doc, headers, rows, widths_cm, center_columns=(0,)):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, value in enumerate(headers):
        fill_cell(table.rows[0].cells[i], value, header=True, center=True, font_size=8.5)
    set_repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        for col_index, value in enumerate(values):
            fill_cell(cells[col_index], str(value), center=col_index in center_columns)
            if row_index % 2 == 1:
                shade(cells[col_index], LIGHT_GRAY)
    set_table_geometry(table, widths_cm)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=100, start=140, bottom=100, end=140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_font(r, size=10, bold=True, color=NAVY)
    r = p.add_run(text)
    set_font(r, size=10)
    set_table_geometry(table, [27.9])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_page_field(paragraph):
    run = paragraph.add_run("Trang ")
    set_font(run, size=9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def build():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(0.9)
    section.right_margin = Cm(0.9)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, before, after in (
        ("Heading 1", 14, 10, 6),
        ("Heading 2", 12.5, 8, 5),
        ("Heading 3", 11.5, 6, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("000000")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("4. KIỂM THỬ HỆ THỐNG")
    set_font(r, size=16, bold=True)

    add_note(
        doc,
        "Lưu ý phạm vi",
        "Tài liệu giữ 14 kịch bản chức năng cốt lõi do người viết lựa chọn. Các kịch bản chưa được xem là đã kiểm thử cho đến khi có kết quả thực tế và bằng chứng. Các tiêu chí phi chức năng là ngưỡng đề xuất, cần điều chỉnh theo môi trường triển khai.",
    )

    add_heading(doc, "4.1. KỊCH BẢN KIỂM THỬ", 1)
    add_heading(doc, "4.1.1. Kiểm thử chức năng", 2)
    add_data_table(
        doc,
        ["Mã kịch bản", "Chức năng", "Mô tả kịch bản kiểm thử", "Dữ liệu đầu vào", "Kết quả mong đợi"],
        FUNCTIONAL,
        [2.0, 2.6, 6.5, 7.3, 9.5],
        center_columns=(0, 1),
    )

    add_heading(doc, "4.1.2. Kiểm thử phi chức năng", 2)
    p = doc.add_paragraph("Các ngưỡng sau là tiêu chí chấp nhận đề xuất. Khi thực thi cần ghi rõ cấu hình máy chủ, dữ liệu, công cụ và số người dùng ảo.")
    p.paragraph_format.keep_with_next = True
    add_data_table(
        doc,
        ["Mã kịch bản", "Thuộc tính", "Mô tả kịch bản kiểm thử", "Dữ liệu/điều kiện", "Kết quả mong đợi"],
        NON_FUNCTIONAL,
        [2.0, 2.8, 7.2, 7.0, 8.9],
        center_columns=(0, 1),
    )

    add_heading(doc, "4.2. KẾT QUẢ THỬ NGHIỆM CÁC KỊCH BẢN", 1)
    add_note(doc, "Nguyên tắc", "Không ghi “Đạt” nếu chưa chạy kịch bản và chưa có bằng chứng như ảnh màn hình, response API, log hoặc báo cáo tải.")
    result_rows = [(code, "Chưa thực hiện", "Chưa đánh giá", "") for code, *_ in FUNCTIONAL]
    add_data_table(
        doc,
        ["Mã kịch bản", "Kết quả thực tế", "Đạt/Không đạt", "Ghi chú (nếu có lỗi)"],
        result_rows,
        [3.0, 9.2, 5.2, 10.5],
        center_columns=(0, 1, 2),
    )

    add_heading(doc, "4.3. XỬ LÝ CÁC TRƯỜNG HỢP NGOẠI LỆ", 1)
    add_data_table(
        doc,
        ["Chức năng", "Tình huống ngoại lệ", "Cách hệ thống xử lý", "Kết quả mong đợi"],
        EXCEPTIONS,
        [4.2, 7.5, 9.0, 7.2],
        center_columns=(0,),
    )

    add_heading(doc, "4.4. QUY ƯỚC ĐÁNH GIÁ", 1)
    definitions = [
        ("Đạt", "Kết quả thực tế thỏa toàn bộ kết quả mong đợi và có bằng chứng."),
        ("Không đạt", "Có ít nhất một điều kiện mong đợi không thỏa; cần ghi bước tái hiện, mức độ và bằng chứng."),
        ("Chưa đánh giá", "Kịch bản chưa được thực thi hoặc chưa đủ bằng chứng để kết luận."),
    ]
    for label, text in definitions:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_font(r, size=11, bold=True)
        r = p.add_run(text)
        set_font(r, size=11)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("BỘ TEST CASE LUẬN VĂN - BALII E-COMMERCE PLATFORM")
    set_font(r, size=8.5, color="666666")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer)

    doc.core_properties.title = "Bộ test case luận văn rút gọn"
    doc.core_properties.subject = "Kiểm thử chức năng, phi chức năng và ngoại lệ"
    doc.core_properties.author = ""
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
