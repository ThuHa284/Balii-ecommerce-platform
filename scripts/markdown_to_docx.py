from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_RE = re.compile(r"^[-*+]\s+(.*)$")
NUMBER_RE = re.compile(r"^\d+\.\s+(.*)$")


@dataclass
class Block:
    kind: str
    lines: List[str]


def parse_blocks(text: str) -> List[Block]:
    lines = text.splitlines()
    blocks: List[Block] = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            fence = stripped
            code_lines: List[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            blocks.append(Block("code", code_lines))
            continue

        if HEADING_RE.match(stripped):
            blocks.append(Block("heading", [stripped]))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if next_line.startswith("|") and "---" in next_line:
                table_lines = [stripped, next_line]
                i += 2
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
                blocks.append(Block("table", table_lines))
                continue

        if BULLET_RE.match(stripped):
            bullet_lines = [stripped]
            i += 1
            while i < len(lines) and BULLET_RE.match(lines[i].strip()):
                bullet_lines.append(lines[i].strip())
                i += 1
            blocks.append(Block("bullet_list", bullet_lines))
            continue

        if NUMBER_RE.match(stripped):
            number_lines = [stripped]
            i += 1
            while i < len(lines) and NUMBER_RE.match(lines[i].strip()):
                number_lines.append(lines[i].strip())
                i += 1
            blocks.append(Block("number_list", number_lines))
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            candidate = lines[i]
            candidate_stripped = candidate.strip()
            if not candidate_stripped:
                break
            if candidate_stripped.startswith("```"):
                break
            if HEADING_RE.match(candidate_stripped):
                break
            if BULLET_RE.match(candidate_stripped):
                break
            if NUMBER_RE.match(candidate_stripped):
                break
            if candidate_stripped.startswith("|") and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if next_line.startswith("|") and "---" in next_line:
                    break
            para_lines.append(candidate_stripped)
            i += 1
        blocks.append(Block("paragraph", para_lines))

    return blocks


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(10.5)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        elem = tbl_cell_mar.find(qn(f"w:{tag}"))
        if elem is None:
            elem = OxmlElement(f"w:{tag}")
            tbl_cell_mar.append(elem)
        elem.set(qn("w:w"), str(value))
        elem.set(qn("w:type"), "dxa")


def apply_run_font(run, font_name: str, font_size: float, color: str | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def add_inline_runs(paragraph, text: str, font_name: str = "Calibri", font_size: float = 11.0) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            apply_run_font(run, font_name, font_size)
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            apply_run_font(run, "Consolas", font_size - 0.5)
        else:
            run = paragraph.add_run(part)
            apply_run_font(run, font_name, font_size)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.start_type = WD_SECTION_START.NEW_PAGE

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def add_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(title)
    apply_run_font(run, "Calibri", 22, "17365D")
    run.bold = True


def render_heading(document: Document, line: str, title_done: bool) -> bool:
    match = HEADING_RE.match(line)
    if not match:
        return title_done
    level = len(match.group(1))
    text = match.group(2).strip()
    if level == 1 and not title_done:
        add_title(document, text)
        return True
    if level > 3:
        level = 3
    document.add_paragraph(text, style=f"Heading {level}")
    return title_done


def render_paragraph(document: Document, lines: Iterable[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    text = " ".join(line.strip() for line in lines)
    add_inline_runs(paragraph, text)


def render_list(document: Document, lines: Iterable[str], numbered: bool) -> None:
    style_name = "List Number" if numbered else "List Bullet"
    for line in lines:
        text = NUMBER_RE.sub(r"\1", line) if numbered else BULLET_RE.sub(r"\1", line)
        paragraph = document.add_paragraph(style=style_name)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
        add_inline_runs(paragraph, text)


def render_code(document: Document, lines: Iterable[str]) -> None:
    for idx, line in enumerate(lines):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0 if idx else 4)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.paragraph_format.right_indent = Inches(0.2)
        run = paragraph.add_run(line if line else " ")
        apply_run_font(run, "Consolas", 9.5)


def normalize_table_line(line: str) -> List[str]:
    trimmed = line.strip().strip("|")
    return [cell.strip() for cell in trimmed.split("|")]


def render_table(document: Document, lines: List[str]) -> None:
    rows = [normalize_table_line(line) for line in lines]
    headers = rows[0]
    data_rows = rows[2:]
    cols = len(headers)
    table = document.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_cell_margins(table)

    widths = compute_column_widths([headers, *data_rows], total_width=6.5)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "E8EEF5")

    for data in data_rows:
        row_cells = table.add_row().cells
        for idx in range(cols):
            value = data[idx] if idx < len(data) else ""
            row_cells[idx].width = Inches(widths[idx])
            set_cell_text(row_cells[idx], value)

    document.add_paragraph()


def compute_column_widths(rows: List[List[str]], total_width: float) -> List[float]:
    col_count = max(len(row) for row in rows)
    weights = [0] * col_count
    for row in rows:
        for idx, cell in enumerate(row):
            weights[idx] = max(weights[idx], max(len(cell), 6))
    total_weight = sum(weights) or col_count
    return [total_width * weight / total_weight for weight in weights]


def build_docx(markdown_path: Path, output_path: Path) -> None:
    text = markdown_path.read_text(encoding="utf-8")
    blocks = parse_blocks(text)
    document = Document()
    configure_document(document)

    title_done = False
    for block in blocks:
        if block.kind == "heading":
            title_done = render_heading(document, block.lines[0], title_done)
        elif block.kind == "paragraph":
            render_paragraph(document, block.lines)
        elif block.kind == "bullet_list":
            render_list(document, block.lines, numbered=False)
        elif block.kind == "number_list":
            render_list(document, block.lines, numbered=True)
        elif block.kind == "code":
            render_code(document, block.lines)
        elif block.kind == "table":
            render_table(document, block.lines)

    footer = document.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_paragraph.add_run("Balii SleepWear Technical Handbook")
    apply_run_font(footer_run, "Calibri", 9)
    footer_run.add_break(WD_BREAK.LINE)
    page_run = footer_paragraph.add_run("Trang ")
    apply_run_font(page_run, "Calibri", 9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_separate)
    page_run._r.append(fld_end)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert simple Markdown to DOCX.")
    parser.add_argument("input_markdown", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()
    build_docx(args.input_markdown, args.output_docx)


if __name__ == "__main__":
    main()
